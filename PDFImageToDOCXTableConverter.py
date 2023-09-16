'''
pip install python-docx Pillow PyMuPDF
'''

#######################
##############

PATH_TO_FOLDER_WITH_PDFS = '/home/user/Documents/SCE_BASE/0_kurs_3_sem/Security_HW/Examen/Lectures'
# PATH_TO_FOLDER_WITH_PDFS = '/home/user/Documents/SCE_BASE/0_kurs_3_sem/Security_HW/Examen/ex_plus_answer'

##############
#######################


import os
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image
import shutil

import fitz  # PyMuPDF

import re

def tryint(s):
    try:
        return int(s)
    except:
        return s

def alphanum_key(s):
    """ Turn a string into a list of string and number chunks.
        "z23a" -> ["z", 23, "a"]
    """
    return [ tryint(c) for c in re.split('([0-9]+)', s) ]

def sort_nicely(l):
    """ Sort the given list in the way that humans expect.
    """
    l.sort(key=alphanum_key)

def set_margins_to_zero(doc):
    for section in doc.sections:
        section.left_margin = 0
        section.right_margin = 0
        section.top_margin = 0
        section.bottom_margin = 0

def choice_size(doc, image_path):
    page_width = doc.sections[0].page_width
    page_height= doc.sections[0].page_height

    img = Image.open(image_path)
    width, height = img.size


    if(width > height):
        return int(page_width / 2)-Inches(0.1), int(page_height/4)
    else:
        return int(page_width / 2.5)+Inches(0.3), int(page_height/2.5)

def create_image_table(doc, image_folder):
    set_margins_to_zero(doc)

    # Get a list of image files in the specified folder
    image_files = [f for f in os.listdir(image_folder) if f.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp'))]

    # Sort the image files by name
    # image_files.sort()
    sort_nicely(image_files)


    # Calculate the available width for the table
    page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    half_page = int(page_width / 2)


    # Create a table with 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.alignment = 0 
    table.autofit = False

    # Set the width of the first column to half of the available width
    table.columns[0].width = half_page
    table.columns[1].width = half_page

    ###
    row = table.add_row().cells
    # Add the image to the first cell
    cell_1 = row[0]
    cell_1.text = os.path.split(image_folder)[1]
    ###


    resize_image_width, resize_image_height = choice_size(doc, os.path.join(image_folder, image_files[0]))


    for i in range(0, len(image_files), 2):
        image_file = image_files[i]
        # Add a new row to the table for each image
        row = table.add_row().cells

        # Add the image to the first cell
        img_path = os.path.join(image_folder, image_file)
        cell_1 = row[0]
        cell_1.paragraphs[0].alignment = None  # Center alignment
        cell_1.add_paragraph().add_run().add_picture(img_path, width=resize_image_width, height=resize_image_height)


        if(i<len(image_files)-1):
            image_file = image_files[i+1]

            # Add the image to the second cell
            img_path = os.path.join(image_folder, image_file)
            cell_2 = row[1]
            cell_2.paragraphs[0].alignment = None # Center alignment
            cell_2.add_paragraph().add_run().add_picture(img_path, width=resize_image_width, height=resize_image_height)
    
def extract_pdf_pages_as_images(pdf_file, output_folder):
    # Create the output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)

    # Open the PDF file
    pdf_document = fitz.open(pdf_file)

    # Iterate through each page in the PDF
    for page_number in range(pdf_document.page_count):
        page = pdf_document.load_page(page_number)
        
        # Convert the page to a pixmap (image)
        pixmap = page.get_pixmap()

        # Save the pixmap as a JPG image
        image_file = os.path.join(output_folder, f"page_{page_number + 1}.jpg")
        pixmap.save(image_file, "jpg")

    # Close the PDF document
    pdf_document.close()

def shrink_lecture_folder(folder):
    pdf_files = [f for f in os.listdir(folder) if f.endswith(('.pdf', ".PDF"))]
    # pdf_files.sort()
    sort_nicely(pdf_files)
    print("pdf_files", pdf_files)

    folder_name = os.path.split(folder)[1]

    output_docx = f"image_table_{folder_name}.docx"
    doc = Document()

    for pdf_file in pdf_files:
        print("pdf_file", pdf_file)
        pdf_file = os.path.join(folder, pdf_file)
        image_foler = pdf_file[:-4]
        extract_pdf_pages_as_images(pdf_file, image_foler)
        create_image_table(doc, image_foler)
        shutil.rmtree(image_foler)


    doc.save(output_docx)

if __name__ == "__main__":
    print("Begin")
    shrink_lecture_folder(PATH_TO_FOLDER_WITH_PDFS)
    print("End")

